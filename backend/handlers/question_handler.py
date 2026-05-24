import asyncio
import traceback
import time
import io
import os
import json
import fitz  # PyMuPDF dùng cho PDF
import docx
from fpdf import FPDF
from schemas.question_schema import GenerateQuizRequest
from models.document_model import get_document_by_id
from models.question_model import get_question_sets, get_quiz_detail_for_export, save_quiz_to_db
from models.activity_model import log_activity
from utils.ai import generate_quiz_from_ai_for_file 
from utils.export_moodle import generate_moodle_xml
from pathlib import Path

# --- CÁC HÀM XỬ LÝ ĐỒNG BỘ (CHẠY NỀN) ---

def _extract_pdf(file_path: str) -> str:
    text = ""
    with fitz.open(file_path) as pdf_doc:
        for page in pdf_doc:
            text += page.get_text() + "\n"
    return text

def _extract_docx(file_path: str) -> str:
    text = ""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def _extract_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

# --- HÀM BẤT ĐỒNG BỘ CHÍNH ĐỂ GỌI TỪ API ---

async def extract_text_from_file(file_path: str, file_ext: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError("Không tìm thấy file trên server")

    ext = file_ext.lower()
    raw_text = ""

    try:
        if ext == '.pdf':
            raw_text = await asyncio.to_thread(_extract_pdf, file_path)
        elif ext in ['.docx', '.doc']:
            raw_text = await asyncio.to_thread(_extract_docx, file_path)
        elif ext == '.txt':
            raw_text = await asyncio.to_thread(_extract_txt, file_path)
        else:
            return f"Hệ thống chưa hỗ trợ trích xuất chữ từ định dạng {ext}"

        clean_text = raw_text.replace('\x00', '').strip()
        return clean_text
    except Exception as e:
        print(f"❌ Lỗi khi đọc file {file_path}: {str(e)}")
        return ""
    
async def generate_quiz_handler(user_id: int, config: GenerateQuizRequest):
    try:
        doc = await get_document_by_id(config.document_id)
        if not doc:
            return {"success": False, "message": "Tài liệu không tồn tại."}
            
        if doc['user_id'] != user_id:
             return {"success": False, "message": "Bạn không có quyền sử dụng tài liệu này."}

        file_path = doc['file_path']
        
        # Bắt đầu đo thời gian
        start_time = time.time()
        ai_result = await generate_quiz_from_ai_for_file(file_path, config)
        # Kết thúc đo thời gian
        gen_time = time.time() - start_time
        
        if isinstance(ai_result, dict) and 'error' in ai_result:
            error_obj = ai_result.get('error') or {}
            error_message = error_obj.get('message', 'Lỗi không xác định từ AI')
            return {"success": False, "message": f"AI Error: {error_message}"}
        
        current_count = len(ai_result.questions)
        target_count = config.total_questions
        if current_count > target_count:
            ai_result.questions = ai_result.questions[:target_count]
        
        # Ghi log hoạt động
        asyncio.create_task(log_activity(
            user_id, 
            "GENERATE", 
            f"Đã tạo bản nháp bộ câu hỏi từ tài liệu {doc['file_name']}"
        ))

        return {
            "success": True,
            "message": "Tạo bộ câu hỏi thành công (Bản nháp)!",
            "data": ai_result.model_dump(),
            "generation_time": gen_time # Gửi thời gian về FE
        }
    except Exception as e:
        traceback.print_exc()
        return {"success": False, "message": f"Hệ thống gặp sự cố: {str(e)}"}

async def save_quiz_final_handler(user_id: int, request_data: dict):
    try:
        from schemas.question_schema import SaveFinalRequest
        req = SaveFinalRequest(**request_data)
        db_result = await save_quiz_to_db(
            user_id, 
            req.quiz_data, 
            req.config, 
            generation_time=req.generation_time,
            accuracy_score=req.accuracy_score
        )
        
        if not db_result["success"]:
            return {"success": False, "message": db_result["message"]}

        # Ghi log hoạt động
        asyncio.create_task(log_activity(
            user_id, 
            "SAVE", 
            f"Đã lưu bộ đề chính thức: {req.quiz_data.title}"
        ))

        return {
            "success": True,
            "message": "Lưu bộ đề chính thức thành công!",
            "data": {
                "set_id": db_result["set_id"],
                "title": req.quiz_data.title,
                "total_questions": len(req.quiz_data.questions)
            }
        }
    except Exception as e:
        traceback.print_exc()
        return {"success": False, "message": f"Lỗi lưu đề thi: {str(e)}"}

async def get_question_sets_handler(user_id: int, search: str = None, document_id: int = None, limit: int = 20, offset: int = 0):
    try:
        result = await get_question_sets(user_id=user_id, search_term=search, document_id=document_id, limit=limit, offset=offset)
        return result
    except Exception as e:
        traceback.print_exc()
        return {"success": False, "message": f"Lỗi hệ thống: {str(e)}"}
    
async def generate_pdf_logic(data):
    try:
        pdf = FPDF()
        pdf.add_page()
        current_file = Path(__file__).resolve()
        project_root = current_file.parent.parent.parent
        font_path = project_root / "fonts" / "Arial.ttf"
        
        if not font_path.exists():
            print(f"❌ Không tìm thấy file font tại: {font_path}")
            return None

        pdf.add_font("ArialUni", "", str(font_path))
        pdf.set_font("ArialUni", size=16)
        pdf.cell(0, 10, txt=str(data.title), new_x="LMARGIN", new_y="NEXT", align="C")
        
        pdf.set_font("ArialUni", size=10)
        pdf.multi_cell(0, 10, txt=str(data.description or ""))
        pdf.ln(5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)

        for i, q in enumerate(data.questions, 1):
            pdf.set_font("ArialUni", size=12)
            cog_level_map = {"RECALL": "Nhận biết", "UNDERSTAND": "Thông hiểu", "APPLY": "Vận dụng"}
            cog_text = cog_level_map.get(q.cognitive_level, q.cognitive_level)
            pdf.multi_cell(0, 10, txt=f"Câu {i} [{cog_text}]: {q.question_text}")
            
            for opt in q.options:
                prefix = "(x)" if opt.is_correct else "( )"
                pdf.set_x(20) 
                pdf.multi_cell(0, 8, txt=f"{prefix} {opt.option_text}")
            pdf.ln(5)

        pdf_content = pdf.output()
        pdf_stream = io.BytesIO(pdf_content)
        pdf_stream.seek(0)
        return pdf_stream
    except Exception as e:
        print(f"❌ Lỗi PDF: {str(e)}")
        traceback.print_exc()
        return None

async def export_quiz_handler(set_id: int, user_id: int, file_format: str):
    try:
        data = await get_quiz_detail_for_export(set_id, user_id)
        if not data:
            return {"success": False, "message": "Không tìm thấy dữ liệu bộ đề"}

        file_stream = None 
        media_type = ""
        file_name = f"Quiz_{set_id}.{file_format}"

        if file_format == "docx":
            doc = docx.Document()
            doc.add_heading(data.title, 0)
            if data.description:
                doc.add_paragraph(data.description)
            
            for i, q in enumerate(data.questions, 1):
                cog_level_map = {"RECALL": "Nhận biết", "UNDERSTAND": "Thông hiểu", "APPLY": "Vận dụng"}
                cog_text = cog_level_map.get(q.cognitive_level, q.cognitive_level)
                doc.add_paragraph(f"Câu {i} [{cog_text}]: {q.question_text}", style='List Number')
                for opt in q.options:
                    prefix = "[x]" if opt.is_correct else "[ ]"
                    doc.add_paragraph(f"    {prefix} {opt.option_text}")
            
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif file_format == "pdf":
            file_stream = await generate_pdf_logic(data)
            if not file_stream:
                return {"success": False, "message": "Lỗi khi tạo luồng PDF"}
            media_type = "application/pdf"

        elif file_format == "moodle":
            questions_list = []
            for q in data.questions:
                questions_list.append({
                    "question_text": q.question_text,
                    "question_type": q.question_type,
                    "cognitive_level": q.cognitive_level,
                    "explanation": q.explanation,
                    "options": [
                        {
                            "option_text": opt.option_text,
                            "is_correct": opt.is_correct,
                            "distractor_logic": opt.distractor_logic
                        } for opt in q.options
                    ]
                })
            
            xml_content = generate_moodle_xml(data.title, questions_list)
            file_stream = io.BytesIO(xml_content.encode("utf-8"))
            file_stream.seek(0)
            media_type = "application/xml"
            file_name = f"Moodle_Quiz_{set_id}.xml"
            
        # Ghi log hoạt động
        asyncio.create_task(log_activity(
            user_id, 
            "EXPORT", 
            f"Đã xuất bộ đề '{data.title}' dưới định dạng {file_format.upper()}"
        ))
            
        return {
            "success": True, 
            "file_stream": file_stream,
            "file_name": file_name, 
            "media_type": media_type
        }
    except Exception as e:
        traceback.print_exc()
        return {"success": False, "message": str(e)}